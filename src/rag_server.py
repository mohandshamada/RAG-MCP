"""
RAG MCP Server - Multi-format document processing with OCR, table extraction
Supports: PDF, Word, Excel, Images
Uses: LangChain + FAISS + Pytesseract + Tabula
"""

import os
import json
import sys
from pathlib import Path
from typing import Optional, List, Dict, Any

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent))

from mcp.server.fastmcp import FastMCP
import fitz  # PyMuPDF
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings  # Updated from deprecated langchain_community
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
import logging
from comparison_engine import ComparisonEngine

# Initialize logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize MCP server
mcp = FastMCP("RAG Multi-Format Document Server")

# Configuration
BASE_DIR = Path(__file__).parent.parent
DATA_DIR = BASE_DIR / "data"
VECTOR_STORE_DIR = DATA_DIR / "vector_stores"
DOCUMENT_CACHE_DIR = DATA_DIR / "document_cache"
METADATA_DIR = DATA_DIR / "metadata"

# Create directories
for dir_path in [DATA_DIR, VECTOR_STORE_DIR, DOCUMENT_CACHE_DIR, METADATA_DIR]:
    dir_path.mkdir(parents=True, exist_ok=True)

# In-memory storage
loaded_documents = {}
vector_stores = {}
documents_metadata = {}

# Initialize comparison engine (pass rag_query function)
comparison_engine = None  # Will be initialized after rag_query is defined


def get_embeddings():
    """Initialize HuggingFace embeddings"""
    return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")


def extract_pdf_content(file_path: str) -> dict:
    """Extract text and metadata from PDF using PyMuPDF"""
    try:
        pdf_content = {"text": "", "pages": [], "tables": [], "images": []}

        # Open PDF with PyMuPDF
        doc = fitz.open(file_path)

        pdf_content["metadata"] = {
            "total_pages": len(doc),
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_type": "pdf"
        }

        # Extract text from each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            pdf_content["pages"].append({
                "page_num": page_num + 1,
                "text": page_text
            })
            pdf_content["text"] += f"\n--- Page {page_num + 1} ---\n{page_text}"

        doc.close()
        return pdf_content
    except Exception as e:
        logger.error(f"Error extracting PDF: {e}")
        raise


def extract_excel_content(file_path: str) -> dict:
    """Extract text and tables from Excel files"""
    try:
        import openpyxl
        
        excel_content = {"text": "", "sheets": [], "tables": []}
        
        workbook = openpyxl.load_workbook(file_path)
        excel_content["metadata"] = {
            "total_sheets": len(workbook.sheetnames),
            "sheet_names": workbook.sheetnames,
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_type": "excel"
        }
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_text = f"\n--- Sheet: {sheet_name} ---\n"
            sheet_data = []
            
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                sheet_data.append(row_data)
                sheet_text += " | ".join(row_data) + "\n"
            
            excel_content["sheets"].append({
                "name": sheet_name,
                "data": sheet_data,
                "text": sheet_text
            })
            excel_content["tables"].append({
                "sheet": sheet_name,
                "data": sheet_data
            })
            excel_content["text"] += sheet_text
        
        return excel_content
    except Exception as e:
        logger.error(f"Error extracting Excel: {e}")
        raise


def extract_word_content(file_path: str) -> dict:
    """Extract text from Word documents"""
    try:
        from docx import Document as DocxDocument
        
        word_content = {"text": "", "paragraphs": [], "tables": []}
        
        doc = DocxDocument(file_path)
        word_content["metadata"] = {
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_type": "word"
        }
        
        # Extract paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text
            word_content["paragraphs"].append({
                "index": para_idx,
                "text": para_text
            })
            word_content["text"] += para_text + "\n"
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            table_text = f"\n--- Table {table_idx + 1} ---\n"
            
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
                table_text += " | ".join(row_data) + "\n"
            
            word_content["tables"].append({
                "index": table_idx,
                "data": table_data,
                "text": table_text
            })
            word_content["text"] += table_text
        
        return word_content
    except Exception as e:
        logger.error(f"Error extracting Word: {e}")
        raise


def extract_image_with_ocr(file_path: str) -> dict:
    """Extract text from images using OCR (Tesseract)"""
    try:
        import pytesseract
        from PIL import Image
        
        image_content = {"text": "", "ocr_data": {}}
        
        img = Image.open(file_path)
        ocr_text = pytesseract.image_to_string(img)
        
        image_content["text"] = ocr_text
        image_content["ocr_data"] = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_type": "image",
            "image_size": img.size,
            "confidence": "OCR extracted"
        }
        image_content["metadata"] = image_content["ocr_data"]
        
        return image_content
    except Exception as e:
        logger.error(f"Error extracting image with OCR: {e}")
        raise


def detect_file_type(file_path: str) -> str:
    """Detect file type from extension"""
    ext = Path(file_path).suffix.lower()
    
    if ext == ".pdf":
        return "pdf"
    elif ext in [".xlsx", ".xls"]:
        return "excel"
    elif ext in [".docx", ".doc"]:
        return "word"
    elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff"]:
        return "image"
    else:
        raise ValueError(f"Unsupported file type: {ext}")


@mcp.tool()
def ingest_document(file_path: str, document_name: str) -> dict:
    """
    Ingest and process document (PDF, Excel, Word, Image with OCR).
    Creates FAISS index and metadata.

    Args:
        file_path: Full path to document
        document_name: Name to identify document

    Returns:
        Ingestion status and metadata

    Raises:
        FileNotFoundError: If file does not exist
        ValueError: If file type is unsupported or file is too large
    """
    try:
        # Validate file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Validate file size (500MB limit from config)
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > 500:
            raise ValueError(f"File size ({file_size_mb:.2f}MB) exceeds maximum allowed size (500MB)")

        file_type = detect_file_type(file_path)
        logger.info(f"Processing {file_type} file: {file_path}")
        
        # Extract content based on file type
        if file_type == "pdf":
            document_content = extract_pdf_content(file_path)
        elif file_type == "excel":
            document_content = extract_excel_content(file_path)
        elif file_type == "word":
            document_content = extract_word_content(file_path)
        elif file_type == "image":
            document_content = extract_image_with_ocr(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Create text chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        
        chunks = text_splitter.split_text(document_content["text"])
        
        # Create Document objects with metadata
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "document_name": document_name,
                    "file_type": file_type,
                    "chunk_id": i,
                    "file_path": file_path,
                    "source": os.path.basename(file_path)
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        # Create FAISS vector store
        embeddings = get_embeddings()
        vector_store = FAISS.from_documents(documents, embeddings)
        
        # Save FAISS index
        store_path = VECTOR_STORE_DIR / document_name
        vector_store.save_local(str(store_path))
        
        # Store metadata
        metadata = {
            "document_name": document_name,
            "file_type": file_type,
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "chunks_created": len(chunks),
            "content_length": len(document_content["text"]),
            **document_content.get("metadata", {})
        }
        
        # Save metadata
        metadata_file = METADATA_DIR / f"{document_name}_metadata.json"
        with open(metadata_file, "w") as f:
            json.dump(metadata, f, indent=2)
        
        # Store in memory
        loaded_documents[document_name] = document_content
        vector_stores[document_name] = {
            "vector_store": vector_store,
            "store_path": str(store_path),
            "num_chunks": len(chunks)
        }
        documents_metadata[document_name] = metadata
        
        return {
            "success": True,
            "document_name": document_name,
            "file_type": file_type,
            "chunks_created": len(chunks),
            "vector_store_path": str(store_path),
            "metadata_saved": str(metadata_file)
        }
    
    except Exception as e:
        logger.error(f"Error ingesting document: {e}")
        return {"success": False, "error": str(e)}


@mcp.tool()
def rag_query(document_name: str, query: str, top_k: int = 5) -> dict:
    """
    Query document using RAG (Retrieval-Augmented Generation).

    Args:
        document_name: Name of indexed document
        query: Query string (natural language)
        top_k: Number of results to return

    Returns:
        Retrieved relevant content with similarity scores

    Raises:
        ValueError: If document is not indexed
    """
    if document_name not in vector_stores:
        raise ValueError(f"Document '{document_name}' is not indexed. Available documents: {list(vector_stores.keys())}")
    
    try:
        vector_store = vector_stores[document_name]["vector_store"]
        
        # Perform semantic search
        results = vector_store.similarity_search_with_score(query, k=top_k)
        
        retrieved_content = {
            "document_name": document_name,
            "query": query,
            "num_results": len(results),
            "results": [
                {
                    "chunk_id": doc.metadata.get("chunk_id", "unknown"),
                    "similarity_score": float(1 - score),  # Convert distance to similarity
                    "content": doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
                    "full_content": doc.page_content
                }
                for doc, score in results
            ]
        }
        
        return retrieved_content
    
    except Exception as e:
        logger.error(f"Error querying document: {e}")
        return {"error": str(e)}


@mcp.tool()
def rag_batch_query(document_names: List[str], query: str, top_k: int = 3) -> dict:
    """
    Query multiple documents using RAG.
    
    Args:
        document_names: List of document names
        query: Query string
        top_k: Results per document
        
    Returns:
        Results from all documents
    """
    results = {}
    
    for doc_name in document_names:
        if doc_name in vector_stores:
            result = rag_query(doc_name, query, top_k)
            results[doc_name] = result
        else:
            results[doc_name] = {"error": f"Document not indexed"}
    
    return {
        "query": query,
        "documents_queried": len(document_names),
        "results_per_document": top_k,
        "findings": results
    }


@mcp.tool()
def extract_tables_from_document(document_name: str) -> dict:
    """
    Extract all tables from a document.
    
    Args:
        document_name: Name of document
        
    Returns:
        All tables found in document
    """
    if document_name not in loaded_documents:
        return {"error": f"Document '{document_name}' not loaded."}
    
    doc_content = loaded_documents[document_name]
    
    return {
        "document_name": document_name,
        "total_tables": len(doc_content.get("tables", [])),
        "tables": doc_content.get("tables", [])
    }


@mcp.tool()
def extract_images_from_document(document_name: str) -> dict:
    """
    Get information about extracted images/OCR data.
    
    Args:
        document_name: Name of document
        
    Returns:
        Image and OCR information
    """
    if document_name not in loaded_documents:
        return {"error": f"Document '{document_name}' not loaded."}
    
    doc_content = loaded_documents[document_name]
    
    return {
        "document_name": document_name,
        "total_images": len(doc_content.get("images", [])),
        "ocr_available": "ocr_data" in doc_content,
        "images": doc_content.get("images", [])
    }


@mcp.tool()
def get_document_summary(document_name: str) -> dict:
    """
    Get summary and metadata for a document.
    
    Args:
        document_name: Name of document
        
    Returns:
        Document metadata and statistics
    """
    if document_name not in documents_metadata:
        return {"error": f"Document '{document_name}' metadata not found."}
    
    metadata = documents_metadata[document_name]
    
    return {
        "document_name": document_name,
        "metadata": metadata,
        "indexed": document_name in vector_stores,
        "vector_store_chunks": vector_stores[document_name].get("num_chunks", 0) if document_name in vector_stores else 0
    }


@mcp.tool()
def list_indexed_documents() -> dict:
    """
    List all indexed documents.
    
    Returns:
        Dictionary of indexed documents
    """
    documents = {}
    
    for doc_name, metadata in documents_metadata.items():
        documents[doc_name] = {
            "file_type": metadata.get("file_type", "unknown"),
            "chunks": vector_stores.get(doc_name, {}).get("num_chunks", 0),
            "content_length": metadata.get("content_length", 0),
            "file_name": metadata.get("file_name", "unknown")
        }
    
    return {
        "total_documents": len(documents),
        "documents": documents
    }


@mcp.tool()
def load_existing_index(document_name: str, store_path: str) -> dict:
    """
    Load previously saved FAISS index.
    
    Args:
        document_name: Name to identify document
        store_path: Path to FAISS store directory
        
    Returns:
        Loading status
    """
    try:
        if not os.path.exists(store_path):
            return {"success": False, "error": f"Store not found: {store_path}"}
        
        embeddings = get_embeddings()
        vector_store = FAISS.load_local(store_path, embeddings)
        
        vector_stores[document_name] = {
            "vector_store": vector_store,
            "store_path": store_path,
            "num_chunks": len(vector_store.docstore._dict) if hasattr(vector_store.docstore, '_dict') else 0
        }
        
        return {
            "success": True,
            "document_name": document_name,
            "store_path": store_path
        }
    
    except Exception as e:
        logger.error(f"Error loading index: {e}")
        return {"success": False, "error": str(e)}


@mcp.tool()
def delete_document_index(document_name: str) -> dict:
    """
    Delete document index and metadata.
    
    Args:
        document_name: Name of document to delete
        
    Returns:
        Deletion status
    """
    try:
        if document_name in vector_stores:
            store_path = vector_stores[document_name]["store_path"]
            del vector_stores[document_name]
            
            # Remove from disk
            if os.path.exists(store_path):
                import shutil
                shutil.rmtree(store_path)
        
        if document_name in documents_metadata:
            del documents_metadata[document_name]
        
        if document_name in loaded_documents:
            del loaded_documents[document_name]
        
        # Remove metadata file
        metadata_file = METADATA_DIR / f"{document_name}_metadata.json"
        if metadata_file.exists():
            metadata_file.unlink()
        
        return {
            "success": True,
            "document_name": document_name,
            "message": f"Index deleted for {document_name}"
        }
    
    except Exception as e:
        logger.error(f"Error deleting index: {e}")
        return {"success": False, "error": str(e)}


@mcp.tool()
def generate_rag_report(document_name: str, queries: List[str]) -> dict:
    """
    Generate comprehensive RAG report with multiple queries.
    
    Args:
        document_name: Document to query
        queries: List of queries to run
        
    Returns:
        Comprehensive report with findings
    """
    if document_name not in vector_stores:
        return {"error": f"Document '{document_name}' not indexed."}
    
    try:
        report = {
            "document_name": document_name,
            "total_queries": len(queries),
            "query_results": []
        }
        
        for query in queries:
            result = rag_query(document_name, query, top_k=3)
            if "error" not in result:
                report["query_results"].append(result)
        
        return report
    
    except Exception as e:
        logger.error(f"Error generating report: {e}")
        return {"error": str(e)}


@mcp.tool()
def compare_document_to_specification(
    document_name: str,
    specifications: List[str],
    spec_name: str = "Specification",
    threshold: float = 0.7
) -> dict:
    """
    Compare document against specifications using RAG.
    
    Args:
        document_name: Name of indexed document
        specifications: List of specification requirements
        spec_name: Name of specification set
        threshold: Similarity threshold (0.0-1.0)
        
    Returns:
        Compliance report with detailed findings
    """
    global comparison_engine
    
    if comparison_engine is None:
        comparison_engine = ComparisonEngine(rag_query)
    
    try:
        result = comparison_engine.compare_document_to_spec(
            document_name,
            specifications,
            spec_name,
            threshold
        )
        
        # Convert to dict for JSON serialization
        return {
            "success": True,
            "document_name": result.document_name,
            "spec_name": result.spec_name,
            "compliance_percentage": result.compliance_percentage,
            "summary": result.summary,
            "statistics": {
                "total_requirements": result.total_requirements,
                "compliant": result.compliant_items,
                "partial": result.partial_items,
                "non_compliant": result.non_compliant_items,
                "unknown": result.unknown_items
            },
            "items": [
                {
                    "requirement_id": item.requirement_id,
                    "requirement_text": item.requirement_text,
                    "status": item.status.value,
                    "found_value": item.found_value,
                    "evidence": item.evidence,
                    "notes": item.notes
                }
                for item in result.items
            ]
        }
    
    except Exception as e:
        logger.error(f"Error comparing document: {e}")
        return {"success": False, "error": str(e)}


@mcp.tool()
def compare_multiple_documents_to_spec(
    document_names: List[str],
    specifications: List[str],
    spec_name: str = "Specification"
) -> dict:
    """
    Compare multiple documents against same specification.
    
    Args:
        document_names: List of document names to compare
        specifications: List of specification requirements
        spec_name: Name of specification set
        
    Returns:
        Comparison results for all documents
    """
    global comparison_engine
    
    if comparison_engine is None:
        comparison_engine = ComparisonEngine(rag_query)
    
    try:
        results = {}
        
        for doc_name in document_names:
            if doc_name not in vector_stores:
                results[doc_name] = {"error": f"Document '{doc_name}' not indexed"}
                continue
            
            result = comparison_engine.compare_document_to_spec(
                doc_name,
                specifications,
                spec_name
            )
            
            results[doc_name] = {
                "document_name": result.document_name,
                "compliance_percentage": result.compliance_percentage,
                "summary": result.summary,
                "statistics": {
                    "total_requirements": result.total_requirements,
                    "compliant": result.compliant_items,
                    "partial": result.partial_items,
                    "non_compliant": result.non_compliant_items,
                    "unknown": result.unknown_items
                }
            }
        
        return {
            "success": True,
            "spec_name": spec_name,
            "documents_compared": len(document_names),
            "results": results
        }
    
    except Exception as e:
        logger.error(f"Error comparing multiple documents: {e}")
        return {"success": False, "error": str(e)}


@mcp.tool()
def generate_compliance_report(
    document_name: str,
    specifications: List[str],
    spec_name: str = "Specification",
    format: str = "text"
) -> dict:
    """
    Generate detailed compliance report in specified format.
    
    Args:
        document_name: Document to report on
        specifications: Specification requirements
        spec_name: Name of specification
        format: Report format ('text', 'json', or 'html')
        
    Returns:
        Formatted compliance report
    """
    global comparison_engine
    
    if comparison_engine is None:
        comparison_engine = ComparisonEngine(rag_query)
    
    try:
        result = comparison_engine.compare_document_to_spec(
            document_name,
            specifications,
            spec_name
        )
        
        report_text = comparison_engine.generate_compliance_report(result, format)
        
        return {
            "success": True,
            "document_name": document_name,
            "spec_name": spec_name,
            "format": format,
            "report": report_text
        }
    
    except Exception as e:
        logger.error(f"Error generating compliance report: {e}")
        return {"success": False, "error": str(e)}


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1 and sys.argv[1] == "dev":
        logger.info("Starting RAG MCP Server in development mode...")
    
    mcp.run()
